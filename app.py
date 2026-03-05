import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import qrcode
from io import BytesIO
import datetime
import pandas as pd

# --- CẤU HÌNH TRANG ---
st.set_page_config(page_title="AI HÀNH CHÍNH CẤP XÃ", layout="wide", page_icon="🏛️")

# --- DATA MẪU 63 TỈNH THÀNH (Demo một số vùng) ---
DATA_HANH_CHINH = {
    "Tuyên Quang": ["Lâm Bình", "Thác Bà", "Sơn Dương", "Chiêm Hóa", "Na Hang"],
    "Hà Nội": ["Quận Ba Đình", "Quận Hoàn Kiếm", "Huyện Đông Anh", "Xã Tiên Dương", "Phường Dịch Vọng"],
    "TP.HCM": ["Quận 1", "Quận 3", "TP. Thủ Đức", "Huyện Cần Giờ", "Xã Bình Chánh"],
    "Đà Nẵng": ["Quận Hải Châu", "Quận Liên Chiểu", "Huyện Hòa Vang"],
    "Cần Thơ": ["Quận Ninh Kiều", "Quận Cái Răng", "Huyện Phong Điền"]
}

LOAI_VAN_BAN = [
    "Công văn (CV)", "Biên bản (BB)", "Quyết định (QĐ)", "Thông báo (TB)", 
    "Chứng nhận (CN)", "Đơn đề nghị", "Giấy xác nhận", "Kế hoạch (KH)", 
    "Báo cáo (BC)", "Phê duyệt", "Triệu tập", "Thông tin"
]

PROMPT_MAU = [
    "Báo cáo công tác tháng 2/2026 xã Lâm Bình, 125 hộ nghèo",
    "Thông báo họp dân ấp Thác Rã, 8h ngày 10/3 về làm đường",
    "Quyết định khen thưởng học sinh giỏi ấp Nà Mức năm 2025",
    "Biên bản họp Mặt trận Tổ quốc ấp Thum Đeng về bầu cử",
    "Chứng nhận hộ nghèo cho hộ ông Nguyễn Văn A, ấp Thác Rã",
    "Kế hoạch phòng chống dịch sốt xuất huyết mùa mưa 2026",
    "Thông báo cưỡng chế thu hồi đất thửa 156 bản đồ địa chính",
    "Giấy xác nhận tạm trú cho bà Lê Thị D từ Hà Nội về xã"
]

# --- HÀM XỬ LÝ DOCX ---
def create_docx(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Header: Quốc hiệu - Tiêu ngữ (Dùng Table để căn chỉnh chuẩn)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    # Cột trái: Tên UBND
    left_cell = table.cell(0, 0).paragraphs[0]
    left_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_cell.add_run(f"UBND {data['xa'].upper()}\n").bold = True
    left_cell.add_run(f"Số: {data['so_hieu']}\n")
    
    # Cột phải: Quốc hiệu
    right_cell = table.cell(0, 1).paragraphs[0]
    right_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_qh = right_cell.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    run_qh.bold = True
    run_tn = right_cell.add_run("Độc lập - Tự do - Hạnh phúc\n")
    run_tn.bold = True
    right_cell.add_run(f"{data['xa']}, ngày {data['ngay'].strftime('%d')} tháng {data['ngay'].strftime('%m')} năm {data['ngay'].strftime('%Y')}")

    doc.add_paragraph("\n")
    
    # Trích yếu
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p_title.add_run(f"V/v: {data['trich_yeu']}")
    title_run.bold = True
    title_run.font.size = Pt(15)

    # Nội dung
    doc.add_paragraph(f"Kính gửi: {data['kinh_gui']}")
    doc.add_paragraph(data['noi_dung'])

    # Chữ ký
    doc.add_paragraph("\n")
    sign_table = doc.add_table(rows=1, cols=2)
    sign_right = sign_table.cell(0, 1).paragraphs[0]
    sign_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_right.add_run("CHỦ TỊCH\n").bold = True
    sign_right.add_run("(Ký, ghi rõ họ tên)\n\n\n\n")
    sign_right.add_run(data['chu_tich']).bold = True

    # Lưu file vào memory
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- GIAO DIỆN SIDEBAR ---
with st.sidebar:
    st.header("🏘️ QUẢN TRỊ CẤP XÃ")
    api_key = st.text_input("OpenAI API Key", type="password")
    
    tinh = st.selectbox("Chọn Tỉnh/Thành", list(DATA_HANH_CHINH.keys()))
    xa = st.selectbox("Chọn Xã/Phường", DATA_HANH_CHINH[tinh])
    
    st.divider()
    chu_tich = st.text_input("Chủ tịch UBND", "Nguyễn Văn A")
    van_thu = st.text_input("Văn thư", "Lê Văn C")
    
    st.info(f"📍 Đang cấu hình cho:\nUBND {xa.upper()} - {tinh.upper()}")

# --- GIAO DIỆN CHÍNH ---
st.title("🏛️ AI HÀNH CHÍNH CẤP XÃ - TRỢ LÝ ĐA NĂNG")
st.caption("Hệ thống soạn thảo văn bản hành chính thông minh dành cho UBND Xã/Phường Việt Nam")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("📋 Thông tin văn bản")
    loai_vb = st.selectbox("Loại văn bản", LOAI_VAN_BAN)
    so_hieu = st.text_input("Số hiệu", f"001/{loai_vb.split('(')[-1].replace(')', '')}-UBND")
    ngay = st.date_input("Ngày ban hành", datetime.date.today())
    kinh_gui = st.text_input("Kính gửi", "UBND Huyện / Nhân dân trên địa bàn xã")
    
    prompt_chon = st.selectbox("🎯 Gợi ý Prompt cấp xã", ["--- Chọn mẫu có sẵn ---"] + PROMPT_MAU)
    user_input = st.text_area("Nội dung/Yêu cầu cụ thể", 
                             value=prompt_chon if prompt_chon != "--- Chọn mẫu có sẵn ---" else "",
                             height=150)

with col2:
    st.subheader("✨ Kết quả AI")
    if st.button("🚀 SOẠN THẢO VĂN BẢN NGAY", type="primary"):
        if not api_key:
            st.error("Vui lòng nhập OpenAI API Key ở thanh bên!")
        else:
            client = OpenAI(api_key=api_key)
            system_prompt = f"""
            BẠN LÀ TRỢ LÝ VĂN THƯ UBND XÃ. Nhiệm vụ: Soạn {loai_vb} theo Nghị định 30/2020/NĐ-CP.
            THÔNG TIN: Xã {xa}, Tỉnh {tinh}, Chủ tịch {chu_tich}.
            CẤU TRÚC 5 PHẦN:
            1. CĂN CỨ: Luật và các quyết định của Huyện/Tỉnh liên quan.
            2. TÌNH HÌNH: Nêu thực trạng tại địa phương (thôn/ấp/số dân).
            3. NỘI DUNG: Các quyết định/thông tin chi tiết (đánh số 1, 2, 3).
            4. TỔ CHỨC THỰC HIỆN: Giao nhiệm vụ cho Công an xã, Địa chính, Thôn trưởng...
            5. KẾT LUẬN: Lời chào và đề nghị phối hợp.
            NGÔN NGỮ: Trang trọng, chính xác, ngắn gọn, dễ hiểu cho người dân.
            """
            
            with st.spinner("AI đang soạn thảo văn bản chuẩn..."):
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Hãy viết nội dung cho văn bản về: {user_input}"}
                    ]
                )
                ai_content = response.choices[0].message.content
                st.session_state['ai_content'] = ai_content
                st.session_state['trich_yeu'] = user_input[:100]

    if 'ai_content' in st.session_state:
        st.markdown(f"### [XEM TRƯỚC VĂN BẢN]")
        st.write(st.session_state['ai_content'])
        
        # Tạo file DOCX
        docx_bytes = create_docx({
            'xa': f"Xã {xa}",
            'tinh': tinh,
            'so_hieu': so_hieu,
            'ngay': ngay,
            'trich_yeu': st.session_state['trich_yeu'],
            'kinh_gui': kinh_gui,
            'noi_dung': st.session_state['ai_content'],
            'chu_tich': chu_tich
        })
        
        st.divider()
        st.download_button(
            label="📥 TẢI FILE WORD (.DOCX) CHUẨN",
            data=docx_bytes,
            file_name=f"{so_hieu.replace('/', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- TÍNH NĂNG PHỤ ---
st.divider()
c1, c2, c3 = st.columns(3)
with c1:
    st.subheader("📱 QR Code Tra cứu")
    qr_data = f"https://vks.xa.{xa.lower().replace(' ', '')}.gov.vn/{so_hieu}"
    img = qrcode.make(qr_data)
    buf = BytesIO()
    img.save(buf)
    st.image(buf.getvalue(), width=150, caption="Mã QR định danh văn bản")

with c2:
    st.subheader("📊 Thống kê xã")
    st.write(f"- Tổng văn bản tháng 3: **12**")
    st.write(f"- Quyết định: **04**")
    st.write(f"- Thông báo: **08**")

with c3:
    st.subheader("💾 Sao lưu")
    st.button("Backup dữ liệu lên Cloud")
    st.button("Xuất Nhật ký văn thư (Excel)")
